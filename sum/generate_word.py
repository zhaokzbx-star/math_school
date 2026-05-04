from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

def generate_word_document():
    doc = Document()
    
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(12)
    
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.5)
    section.right_margin = Inches(1.5)
    
    # 第一页：封面
    cover_path = r'd:\MyProject\Math_school\sum\封面.pdf'
    if os.path.exists(cover_path):
        try:
            import fitz
            pdf_document = fitz.open(cover_path)
            page = pdf_document.load_page(0)
            pix = page.get_pixmap()
            cover_image_path = r'd:\MyProject\Math_school\sum\封面.png'
            pix.save(cover_image_path)
            
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run()
            run.add_picture(cover_image_path, width=Inches(6))
        except ImportError:
            paragraph = doc.add_paragraph('2026吉林大学数学建模竞赛封面')
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.style.font.size = Pt(16)
            paragraph.style.font.bold = True
    else:
        paragraph = doc.add_paragraph('2026吉林大学数学建模竞赛封面')
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.style.font.size = Pt(16)
        paragraph.style.font.bold = True
    
    doc.add_page_break()
    
    # 第二页：题目和摘要（开始页码）
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    
    # 题目
    title = doc.add_heading('城市极端降雨下内涝韧性提升与疏散路径优化', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.style.font.name = '黑体'
    title.style.font.size = Pt(16)
    
    doc.add_paragraph()
    
    # 摘要
    abstract_heading = doc.add_heading('摘要', level=1)
    abstract_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    abstract_heading.style.font.name = '黑体'
    abstract_heading.style.font.size = Pt(14)
    
    abstract_text = """本文针对北方城市短时极端降雨引发的内涝灾害，以A市某3.2 km²中心城区片区一次历时1小时、累计雨量65 mm的暴雨事件为研究对象，围绕积水动态演化、内涝韧性评价、动态疏散优化与韧性提升改造四个核心问题，构建了"物理机理—系统评价—动态决策—工程优化"的完整数学模型体系，为城市内涝智慧治理提供定量化决策支撑。
    
    问题一，基于水量平衡原理，构建引入管网排水折减系数Cd与街区汇水乘数θ的一维离散时间积水动态演化模型，刻画极端降雨下管网排水能力衰减与下垫面汇流放大效应；通过三次样条插值实现降雨序列逐分钟细化，并利用分段函数建立积水深度与通行能力系数、安全度的映射关系。采用最小二乘法反演标定参数，得到Cd = 0.100、θ = 12.05，模型拟合决定系数R² = 0.947，均方根误差仅5.631 cm，预测精度满足应急决策需求。结果显示，L5路段积水峰值达36.4 cm，60分钟时通行能力完全丧失，为片区内涝最高风险节点。
    
    问题二，从抵御、恢复、适应三个维度构建包含9项指标的内涝韧性评价体系，提出MC-EWM-TOPSIS综合评价模型。通过蒙特卡洛模拟生成500组指标样本以量化输入不确定性，采用熵权法客观确定指标权重，测算得到研究片区现状综合韧性得分为0.4308，处于中等偏弱水平。敏感性分析表明，绿色调蓄容积（敏感度0.1978）是制约系统韧性的核心因素，预警覆盖率与抢修覆盖效率次之。基于加权TOPSIS的物理脆弱性指数识别出L5、L2、L8为高脆弱性路段，为靶向治理与疏散规避提供精准依据。
    
    问题三，针对传统静态路径规划在动态积水场景下完全失效的痛点，建立以总疏散时间最短与路径安全度最高为双目标的动态时空规划模型。将时间维度嵌入路网拓扑，设计改进的时变Dijkstra算法，依据实时积水深度动态更新路段通行阻抗，并引入断路剪枝策略提升搜索效率。仿真结果表明，最优动态疏散方案路径为"起点→B→D→终点"，总耗时4.62分钟，综合安全度0.3240；传统静态最短路径因高风险路段断路，耗时14.25分钟，安全度仅0.0300。动态规划方案较静态导航疏散效率提升68%，安全度提升10.8倍，有效验证了本算法在规避内涝断路风险与保障疏散安全方面的显著优越性。
    
    问题四，在5000万元预算约束下，以管网扩容长度、调蓄池数量、透水铺装面积与预警系统建设为决策变量，建立以总成本最低、韧性增益最大、退水时间最短为目标的多目标混合整数优化模型。通过蒙特卡洛采样生成15000个候选方案，采用帕累托支配关系提取非劣解集，并结合K-Means聚类归纳出三类典型改造方案：保守防御型（成本1226万元，韧性增益0.097，退水时间154.9 min）、均衡拐点型（成本2984万元，韧性增益0.171，退水时间125.5 min）与激进海绵型（成本4569万元，韧性增益0.265，退水时间106.9 min）。鲁棒性测试表明，在降雨强度增强15%、设施排水效能下降20%的复合极端扰动下，各方案退水延迟率均低于55%，方案具有良好的工程可靠性与场景适应性。
    
    本文构建的"积水预测→韧性诊断→路径决策→改造优化"模型体系，实现了从状态刻画到工程决策的完整逻辑闭环，可为城市内涝实时风险监测、应急指挥调度及海绵城市基础设施改造提供科学的定量分析工具与决策支持。"""
    
    abstract_paragraph = doc.add_paragraph(abstract_text)
    abstract_paragraph.style.font.name = '宋体'
    abstract_paragraph.style.font.size = Pt(12)
    abstract_paragraph.line_spacing = 1.0
    
    keywords = doc.add_paragraph('关键词：城市内涝；积水动态演化模型；内涝韧性评价；MC-EWM-TOPSIS；时变Dijkstra算法；动态疏散路径规划；多目标优化；帕累托前沿')
    keywords.style.font.name = '宋体'
    keywords.style.font.size = Pt(12)
    
    doc.add_page_break()
    
    # 第三页开始：正文
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    
    # 一、问题重述
    heading1 = doc.add_heading('一、问题重述', level=1)
    heading1.style.font.name = '黑体'
    heading1.style.font.size = Pt(14)
    heading1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    content1 = """近年来，北方城市短时极端强降雨频次显著上升，引发的道路积水与交通中断已成为城市公共安全的重要威胁。本文以A市某中心城区片区为研究对象，针对一次历时1小时、累计雨量65 mm的极端降雨事件，基于给定数据建立数学模型，研究以下四个问题：
    
    1. 积水动态演化建模：基于附表1的路网基础属性、附表2的降雨时序数据、附表3的通行能力映射规则以及附表4的逐时段积水观测数据，建立路段积水深度动态演化模型，预测各路段在降雨历时内不同时刻的积水深度，并据此确定通行能力系数与安全度评分，验证模型的合理性与适用性。
    
    2. 内涝韧性综合评价：依据附表5给出的抵御、恢复、适应三维度韧性评价指标体系，建立片区内涝韧性综合评价模型，计算现状综合韧性得分，分析各指标对整体韧性的贡献程度，识别系统薄弱环节与薄弱路段。
    
    3. 动态疏散路径优化模型：以总疏散时间最短与路径安全度最高为双目标，结合问题一得到的积水动态演化结果与附表3的通行能力约束，建立动态疏散路径优化模型，设计可实时更新的路径搜索算法，并给出典型起点到终点的最优动态路径及对应的时间与安全度指标。
    
    4. 内涝韧性提升综合改造方案：依据附表6给出的各项改造措施成本与效果数据，在有限预算约束下，建立以总成本最低、韧性提升最大、消退时间最短为目标的多目标优化模型，确定最优措施组合，输出可落地的综合改造方案，并对其进行效果评估与鲁棒性分析。"""
    
    doc.add_paragraph(content1)
    
    # 二、问题分析
    heading2 = doc.add_heading('二、问题分析', level=1)
    heading2.style.font.name = '黑体'
    heading2.style.font.size = Pt(14)
    heading2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    content2 = """### 2.1 积水动态演化建模
    
    本问题核心在于基于有限观测数据建立准确的积水深度预测模型，属于水动力学与参数估计范畴。通过水量平衡原理构建水动力学方程，利用实测数据标定模型参数，验证精度并分析误差来源。
    
    ### 2.2 内涝韧性综合评价
    
    本问题旨在科学量化城市内涝韧性并识别薄弱环节，涉及多准则决策分析与敏感性分析方法。通过构建多层次评价指标体系，采用客观赋权法确定权重，计算综合韧性得分并定位薄弱环节。
    
    ### 2.3 动态疏散路径优化模型
    
    本问题需在内涝动态变化场景下实现时间最短与安全度最高的双重目标，涉及动态规划与图论方法。通过构建时空网络模型，设计动态路径搜索算法，对比静态与动态方案的性能差异。
    
    ### 2.4 内涝韧性提升综合改造方案
    
    本问题在有限预算约束下寻求改造措施的最优组合，实现成本、韧性、效率的综合最优，涉及多目标优化与帕累托最优理论。通过建立多目标优化模型，采用NSGA-II算法求解帕累托前沿，提取典型方案并进行鲁棒性分析。"""
    
    doc.add_paragraph(content2)
    
    # 五、模型的建立与求解
    heading5 = doc.add_heading('五、模型的建立与求解', level=1)
    heading5.style.font.name = '黑体'
    heading5.style.font.size = Pt(14)
    heading5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 5.1 问题一
    heading5_1 = doc.add_heading('5.1 问题一：城市路网积水动态演化模型', level=2)
    heading5_1.style.font.name = '黑体'
    heading5_1.style.font.size = Pt(12)
    
    content5_1 = """城市内涝的本质是极端降雨下"汇水-排涝"物理过程的时空失衡。本问要求明确在60分钟极端暴雨历时内，8条核心路段积水深度的动态演化过程。传统的三维水动力学漫溢模型计算量巨大且极度依赖精细的高程与管网拓扑数据，不适用于应急抢险阶段的快速推演。因此，本文基于水量平衡原理，引入"地形坡度放大因子"与基于标高落差的"管网顶托失效机制"，构建了一维离散时间积水动态演化模型。
    
    #### 5.1.1 数据处理
    
    原始降雨数据以5分钟为时间间隔给出，为获得连续时间尺度下的降雨输入，采用三次样条插值方法对降雨序列进行平滑处理，将其转换为逐分钟降雨强度。
    
    #### 5.1.2 模型建立
    
    基于地表水量平衡与一维管网排水机理，建立路段的积水深度动态演化方程：
    
    ∂h/∂t = I_rain - I_infiltration - Q_drain/A
    
    其中降雨强度I_rain为单位时间单位汇水面积上的降雨体积，下渗率I_infiltration为单位时间单位面积的下渗体积，Q_drain为管网排水流量，A为修正后的实际汇水面积。
    
    排水流量Q_drain由管网排水折减系数Cd乘以设计排水能力Q_design得到：Q_drain = Cd · Q_design
    
    汇水面积A由路段长度、宽度及汇水乘数确定：A = θ · L · W
    
    通行能力系数C与安全度S根据积水深度通过分段函数确定，反映积水对交通的不同影响程度。
    
    #### 5.1.3 模型求解
    
    采用有限差分法进行离散求解，参数标定采用最小二乘法，目标函数为预测值与实测值的均方误差。
    
    #### 5.1.4 结果分析
    
    通过最小二乘优化标定得到模型参数：管网排水折减系数Cd = 0.100，街区汇水乘数θ = 12.05。Cd反映实际排水能力相对于理论设计值的折减程度，取值0.100表明现状管网排水效率较低；θ表征街区汇水系统的放大效应，12.05倍的放大效应符合城市下垫面复杂、汇流时间缩短的物理规律。"""
    
    doc.add_paragraph(content5_1)
    
    # 图5.1.4-1
    para = doc.add_paragraph('图5.1.4-1 各路段积水深度动态演化图')
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_picture(r'd:\MyProject\Math_school\a1\answer1图一_积水深度动态演化.png', width=Inches(5.5))
    
    content5_1_2 = """积水深度动态演化图展示了各路段积水深度随时间变化曲线。结果显示：L5路段积水最为严重，峰值达36.4 cm（60分钟时），通行能力降至0；L3、L7路段积水相对较轻，峰值分别为11.9 cm和15.3 cm。各路段积水均呈先升后降趋势，L1、L2、L4、L8在45分钟后趋于稳定。"""
    
    doc.add_paragraph(content5_1_2)
    
    # 图5.1.4-2
    para = doc.add_paragraph('图5.1.4-2 模型预测值与实测值拟合散点图')
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_picture(r'd:\MyProject\Math_school\a1\answer1图二_拟合散点图.png', width=Inches(5.5))
    
    content5_1_3 = """拟合散点图展示了模型预测值与实测值的对比，计算得到R² = 0.947，RMSE = 5.631 cm，MAE = 3.872 cm。R²接近1表明模型能够解释94.7%的实测值变异，预测精度较高。"""
    
    doc.add_paragraph(content5_1_3)
    
    # 图5.1.4-3
    para = doc.add_paragraph('图5.1.4-3 各路段通行能力热力图')
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_picture(r'd:\MyProject\Math_school\a1\answer1图三_通行能力热力图.png', width=Inches(5.5))
    
    content5_1_4 = """通行能力热力图呈现了各路段在不同时刻的通行能力分布，L5路段在45分钟后完全中断，与积水演化规律一致。综合结论：本模型通过物理机理与数据驱动相结合的方式，成功构建了城市内涝积水预测模型。L5路段识别为排水薄弱环节，为后续韧性提升规划提供依据。"""
    
    doc.add_paragraph(content5_1_4)
    
    # 5.2 问题二
    heading5_2 = doc.add_heading('5.2 问题二：内涝韧性综合评价与诊断模型', level=2)
    heading5_2.style.font.name = '黑体'
    heading5_2.style.font.size = Pt(12)
    
    content5_2 = """#### 5.2.1 解题思路与模型选择
    
    本问题旨在建立内涝韧性综合评价模型，采用MC-EWM-TOPSIS算法进行系统层评价。MC（蒙特卡洛模拟）用于处理指标不确定性，EWM（熵权法）客观计算权重，TOPSIS基于理想解相似度计算综合得分；同时采用多准则路段脆弱性诊断方法识别薄弱路段。
    
    #### 5.2.2 数据处理
    
    本问题数据来源于附表5，包含9个二级指标的现状值、理想值与最差值，涵盖抵御能力、恢复能力、适应能力三个维度。预处理阶段采用截断正态分布对每个指标进行蒙特卡洛模拟，生成500次随机采样，充分考虑指标取值的不确定性。
    
    #### 5.2.3 模型建立
    
    **系统层评价模型（MC-EWM-TOPSIS）**：
    
    步骤1-极向标准化：对于正向指标和负向指标分别采用不同的标准化公式。
    
    步骤2-熵权法计算权重：通过信息熵计算各指标的客观权重。
    
    步骤3-TOPSIS综合得分：计算各方案与理想解和负理想解的距离，得到相对贴近度作为综合得分。
    
    **路段层脆弱性诊断模型**：综合考虑路段五个维度的物理特性，采用Min-Max标准化后计算加权TOPSIS贴近度。
    
    **薄弱环节识别方法**：采用扰动分析法识别系统关键影响因素。"""
    
    doc.add_paragraph(content5_2)
    
    # 图5.2.4-1
    para = doc.add_paragraph('图5.2.4-1 主观权重与MC-EWM客观权重对比图')
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_picture(r'd:\MyProject\Math_school\a2\answer2图一_权重对比.png', width=Inches(5.5))
    
    # 图5.2.4-2
    para = doc.add_paragraph('图5.2.4-2 系统韧性诊断雷达图')
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_picture(r'd:\MyProject\Math_school\a2\answer2图二_韧性诊断雷达图.png', width=Inches(5.5))
    
    content5_2_2 = """基于MC-EWM-TOPSIS模型计算得到研究片区综合韧性评分为0.4308，处于中度脆弱状态。主观权重与MC-EWM客观权重存在差异，验证了客观赋权的必要性。韧性诊断雷达图直观展示了系统在三个维度的韧性表现，识别出适应能力为突出短板。"""
    
    doc.add_paragraph(content5_2_2)
    
    # 图5.2.4-4
    para = doc.add_paragraph('图5.2.4-4 各路段脆弱性指数分布柱状图')
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_picture(r'd:\MyProject\Math_school\a2\answer2图四_薄弱路段识别.png', width=Inches(5.5))
    
    content5_2_3 = """基于路段层脆弱性模型计算的物理脆弱性指数显示，L5路段脆弱性最高（0.7631），主要特征为最低标高且最大积水深度达36.4cm。L2和L8路段紧随其后，呈现出低洼地形与排水能力不足的共性问题。"""
    
    doc.add_paragraph(content5_2_3)
    
    # 图5.2.4-3
    para = doc.add_paragraph('图5.2.4-3 各指标敏感度系数排名图')
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_picture(r'd:\MyProject\Math_school\a2\answer2图三_薄弱环节识别.png', width=Inches(5.5))
    
    content5_2_4 = """通过扰动分析法对9项指标进行敏感性分析，绿色调蓄容积的敏感度系数高达0.1978，远超其他指标，说明其对系统韧性具有决定性影响。其次为预警覆盖率和抢修覆盖效率，均属于恢复与适应能力维度。问题二从系统层和路段层两个维度完成了韧性评估与薄弱环节识别。"""
    
    doc.add_paragraph(content5_2_4)
    
    # 5.3 问题三
    heading5_3 = doc.add_heading('5.3 问题三：城市内涝应急疏散路径优化', level=2)
    heading5_3.style.font.name = '黑体'
    heading5_3.style.font.size = Pt(12)
    
    content5_3 = """#### 5.3.1 解题思路与模型选择
    
    本问题旨在建立动态疏散路径优化模型，通过动态时空规划算法实现总疏散时间最短与路径安全度最高的双目标优化。该算法将时间维度纳入路径规划，能够根据实时积水数据动态更新路径。
    
    #### 5.3.2 数据处理
    
    将路网拓扑转换为无向图，通过线性插值实现任意时刻的积水深度估计，根据积水深度映射为通行能力系数和安全度。
    
    #### 5.3.3 模型建立
    
    本问题以疏散时间最短和路径安全度最高为双目标，采用加权线性和法将双目标转化为单目标优化，构建时空路径成本函数。
    
    #### 5.3.4 模型求解与算法设计
    
    采用时变Dijkstra算法，构建时空网络，实时提取最小成本路径进行动态扩展，引入剪枝策略避免无效搜索。"""
    
    doc.add_paragraph(content5_3)
    
    # 图5.3.5-1
    para = doc.add_paragraph('图5.3.5-1 路网拓扑与最优疏散路径图')
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_picture(r'd:\MyProject\Math_school\a3\answer3图二_路网拓扑与最优路径.png', width=Inches(5.5))
    
    # 图5.3.5-2
    para = doc.add_paragraph('图5.3.5-2 各路段积水深度动态演变图')
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_picture(r'd:\MyProject\Math_school\a3\answer3图一_积水深度动态演变.png', width=Inches(5.5))
    
    # 图5.3.5-3
    para = doc.add_paragraph('图5.3.5-3 动态与静态疏散方案对比图')
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_picture(r'd:\MyProject\Math_school\a3\answer3图三_疏散方案对比.png', width=Inches(5.5))
    
    content5_3_2 = """基于时变Dijkstra算法计算得到最优疏散路径为"起点→B→D→终点"，总耗时4.62分钟，综合安全度0.3240。静态导航方案选择最短路径"起点→A→终点"，但因L5路段积水深度超过30cm导致断路，最终耗时14.25分钟，安全度仅0.0300。动态时空规划算法相比静态导航方案，耗时缩短约68%，安全度提升10.8倍。"""
    
    doc.add_paragraph(content5_3_2)
    
    # 5.4 问题四
    heading5_4 = doc.add_heading('5.4 问题四：城市内涝韧性提升的多目标规划', level=2)
    heading5_4.style.font.name = '黑体'
    heading5_4.style.font.size = Pt(12)
    
    content5_4 = """#### 5.4.1 解题思路
    
    本问题旨在5000万元预算约束下，通过优化管网扩容长度、调蓄池数量、透水铺装面积和预警系统建设四个决策变量，实现成本最低、韧性增益最大、退水时间最短三个目标的协同优化。采用蒙特卡洛采样结合帕累托支配排序的方法提取帕累托前沿，再通过K-Means聚类获取三类典型方案。
    
    #### 5.4.2 数据处理
    
    以附表6给出的四项改造措施的成本、效果与物理边界为输入数据，采用均匀分布对决策变量进行大规模随机采样，生成15,000个初始解，筛选满足预算约束的可行解。
    
    #### 5.4.3 模型建立
    
    建立兼顾总成本最低、韧性增益最大、退水时间最短的多目标优化模型。
    
    #### 5.4.4 模型求解与算法设计
    
    采用"蒙特卡洛采样—帕累托筛选—K-Means聚类"三阶段算法，提取三个典型改造方案：保守防御型、均衡拐点型和激进海绵型。"""
    
    doc.add_paragraph(content5_4)
    
    # 图5.4.5-1
    para = doc.add_paragraph('图5.4.5-1 多目标优化帕累托前沿图')
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_picture(r'd:\MyProject\Math_school\a4\answer4图一_帕累托前沿.png', width=Inches(5.5))
    
    # 图5.4.5-2
    para = doc.add_paragraph('图5.4.5-2 成本效益曲线与性价比拐点图')
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_picture(r'd:\MyProject\Math_school\a4\answer4图二_成本效益曲线.png', width=Inches(5.5))
    
    content5_4_2 = """基于帕累托前沿提取和K-Means聚类，得到三个典型改造方案：
    
    | 方案类型 | 总改造成本(万元) | 韧性综合增益 | 正常退水时间(min) |
    |---------|----------------|-------------|------------------|
    | 保守防御型 | 1225.72 | 0.0968 | 154.9 |
    | 均衡拐点型 | 2984.26 | 0.1706 | 125.5 |
    | 激进海绵型 | 4568.66 | 0.2652 | 106.9 |
    
    保守防御型方案以最低成本实现基础韧性提升，均衡拐点型方案在成本与效益之间取得最优平衡，激进海绵型方案以接近预算上限的成本换取最高韧性增益。"""
    
    doc.add_paragraph(content5_4_2)
    
    # 结论
    heading6 = doc.add_heading('六、结论', level=1)
    heading6.style.font.name = '黑体'
    heading6.style.font.size = Pt(14)
    heading6.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    content6 = """### 6.1 主要结论
    
    本研究通过物理机理建模、数据驱动优化与多目标决策方法，构建了城市内涝韧性分析模型体系。积水预测模型标定得到管网排水折减系数Cd = 0.100、街区汇水乘数θ = 12.05，模型R² = 0.947表明预测精度较高，可准确捕捉积水演化规律。系统韧性综合评分为0.4308，处于中等水平，其中绿色调蓄容积是影响韧性的最关键因素。动态时空规划算法相比静态导航方案，疏散效率提升68%、安全度提升10.8倍。多目标优化成功生成三个典型方案：保守防御型（成本1225万元）、均衡拐点型（成本2984万元）和激进海绵型（成本4569万元），分别适用于预算有限、追求平衡和最高韧性等不同场景。
    
    ### 6.2 应用价值与现实意义
    
    本研究建立的模型体系可为城市内涝治理提供科学决策依据。动态路径规划算法可集成到城市应急指挥系统，实现实时疏散指挥与最优路径推荐。韧性评价结果可指导城市基础设施规划，优先改造L5等排水薄弱路段。三个典型改造方案可为政府部门在不同预算约束下提供差异化选择，实现有限资金的最大效益。"""
    
    doc.add_paragraph(content6)
    
    # 参考文献
    heading7 = doc.add_heading('七、参考文献', level=1)
    heading7.style.font.name = '黑体'
    heading7.style.font.size = Pt(14)
    heading7.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    references = """[1] Zhang, M., et al. "A New Urban Waterlogging Simulation Method Based on Multi-Factor Correlation." Water, vol. 14, no. 9, 2022, p. 1421.
    
    [2] Yang, J., et al. "Modelling Trends in Urban Flood Resilience towards Improving the Adaptability of Cities." Water, vol. 16, no. 11, 2024, p. 1614.
    
    [3] Wang, X., et al. "Optimal Evacuation Route Planning of Urban Personnel at Different Risk Levels of Flood Disasters Based on the Improved 3D Dijkstra's Algorithm." Sustainability, vol. 14, no. 16, 2022, p. 10250.
    
    [4] Shen, L., et al. "Building a multi-objective optimization model for Sponge City projects." Sustainable Cities and Society, vol. 83, 2022, p. 103946."""
    
    doc.add_paragraph(references)
    
    # 保存文档
    doc.save(r'd:\MyProject\Math_school\sum\论文.docx')
    print("Word文档已生成：d:\\MyProject\\Math_school\\sum\\论文.docx")

if __name__ == '__main__':
    generate_word_document()